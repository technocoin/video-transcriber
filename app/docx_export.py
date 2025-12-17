from typing import List, Dict
from pathlib import Path
from docx import Document


def _find_placeholder_paragraph(doc: Document, placeholder: str):
    for p in doc.paragraphs:
        if placeholder in p.text:
            return p
    return None


def rows_to_docx(rows: List[Dict], out_path: str, template_path: str):
    template_file = Path(template_path)
    if not template_file.exists():
        # fallback: blank doc if template missing
        doc = Document()
        doc.add_heading("Video Transcript", level=1)
        _add_table(doc, rows)
        doc.save(out_path)
        return

    doc = Document(str(template_file))

    placeholder = _find_placeholder_paragraph(doc, "{{TABLE_HERE}}")
    if placeholder is None:
        # fallback: append at end
        doc.add_paragraph("")
        _add_table(doc, rows)
        doc.save(out_path)
        return

    # clear placeholder text
    placeholder.text = ""

    # Insert table directly after placeholder paragraph
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Sound"
    hdr[2].text = "Vision"
    hdr[3].text = "Compliance"

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.get("Time", ""))
        row_cells[1].text = str(r.get("Sound", ""))
        row_cells[2].text = str(r.get("Vision", ""))
        row_cells[3].text = str(r.get("Compliance", ""))

    doc.save(out_path)


def _add_table(doc: Document, rows: List[Dict]):
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Sound"
    hdr[2].text = "Vision"
    hdr[3].text = "Compliance"

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.get("Time", ""))
        row_cells[1].text = str(r.get("Sound", ""))
        row_cells[2].text = str(r.get("Vision", ""))
        row_cells[3].text = str(r.get("Compliance", ""))

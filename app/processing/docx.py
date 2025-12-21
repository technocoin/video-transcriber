from docx import Document
from docx.shared import Pt
from pathlib import Path


def generate_docx(
    video_name: str,
    transcript: list[dict],
    frames: list[str],
    output_dir: str,
) -> str:
    """
    Generate a DOCX transcript file with timestamps.
    Returns the path to the created DOCX file.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / f"{video_name}.docx"

    doc = Document()

    # -----------------------------
    # Title
    # -----------------------------
    title = doc.add_heading(video_name, level=1)
    title.alignment = 1  # center

    doc.add_paragraph("")

    # -----------------------------
    # Transcript body
    # -----------------------------
    for seg in transcript:
        start = _format_time(seg.get("start", 0))
        end = _format_time(seg.get("end", 0))
        text = seg.get("text", "").strip()

        if not text:
            continue

        p = doc.add_paragraph()

        ts = p.add_run(f"[{start} → {end}] ")
        ts.bold = True
        ts.font.size = Pt(10)

        content = p.add_run(text)
        content.font.size = Pt(11)

        doc.add_paragraph("")

    doc.save(output_path)

    return str(output_path)


def _format_time(seconds: float) -> str:
    minutes = int(seconds // 60)
    seconds = int(seconds % 60)
    return f"{minutes:02d}:{seconds:02d}"
